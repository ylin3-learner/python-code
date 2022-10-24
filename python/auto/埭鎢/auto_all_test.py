# coding:utf-8

'''
现在将每个人工资数据发送至员工本人，要求如下
1、使用Python的xlrd模块读取Excel中的数据
2、为每个人制作工资单文件，文件名为”员工姓名.docx”， # 要先有全局樣式 -from docx.enum.style import WD_STYLE_TYPE
3、将工资单文件作为附件，发送至员工本人的邮箱
'''

import xlrd

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT


excel = xlrd.open_workbook('salary.xlsx')
worksheet = excel.sheet_by_name('研发部工资表')


all_data = []

doc = Document()
# 创建函数，读取excel表格中的内容
def read_excel():
    for i in worksheet.get_rows():
        content = []
        for j in i:
            content.append(j.value)
        all_data.append(content)
    return all_data


title = all_data[0]

print(all_data)


def word_style():
    # 设置字体和大小和格式
    style = doc.styles['Normal']  # 全局樣式
    style.font.size = Pt(11)
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 使中文字也能微軟雅黑

    # 设置标题格式和文字大小
    h = doc.add_heading('', 0)  # 第二個參數為數字, 來標示幾級標題的
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 局中
    _h = h.add_run('5月工资单')
    _h.font.name = '微软雅黑'
    _h.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    _h.bold = True

    print(title, type(title), '---')
# title = ['姓名', '性别', '部门', '基本工资', '绩效工资', '餐补', '社保', '公积金', '代扣个税', '实发工资', '邮箱']

# 定义工资单邮件函数
def word_salary_email(all_data):
    word_style()  # 直接在函數中自動調用格式函數
    table = doc.add_table(rows=1, cols=worksheet.ncols, style='Light List Accent 5')

    # 增加表格，這是表格頭
    title_cells = table.rows[0].cells  # 表格列對象

    for i in range(len(title)):
        title_cells[i].text = title[i]
'''
    Basic cell access protocol
    There are three ways to access a table cell:

    Table.cell(row_idx, col_idx)
    Row.cells[col_idx]
    Column.cells[col_idx]
    
table = document.add_table(3, 3)
middle_cell = table.cell(1, 1)
table.rows[1].cells[1] == middle_cell
>> >True
 
table.columns[1].cells[1] == middle_cell
>> >True
'''

# 如何使用雙循環填入資料??
    # 填充单元格
    for data in all_data[1:]:  # 利用索引避開第一個元素
    # add_row:在原有表格的基础上增加一行数据
        row_cells = table.rows[1].cells
        for index, content in enumerate(data[0:-1]):
            # 数据类型转换，浮点数不能直接写入单元格，对数据进行转换
            content_to_str = str(content)
            row_cells[index].text = content_to_str
        # 不知道如何使文件命名為員工名稱.docx?
        doc.save("{}.docx".format(data[0])) # name
# data = [小慕	男	技术部	15000.0	1800.0	500.0	1713.6	2016.0	635.94	12934.46	xiaomu@imooc.com]
'''
for name, sex, part, basic_money, extra_money, food_cost, society, public_fund, fee, real_money, email in all_data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = part
    row_cells[3].text = str(basic_money)
    row_cells[4].text = str(extra_money)
    row_cells[5].text = str(food_cost)
    row_cells[6].text = str(society)
    row_cells[7].text = str(public_fund)
    row_cells[8].text = str(fee)
    row_cells[9].text = str(real_money)
    row_cells[10].text = email
'''
'''
records = (
    (1, '張三', '電工'),
    (2, '張五', '老闆'),
    (3, '馬六', 'IT'),
    (4, '李四', '工程師')
)

# 遍歷數據並展示
for id, name, work in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = name
    row_cells[2].text = work
'''

# 不知道如何使用gmail發送郵件?

# 第三方的smtp
def send_mail(mail_user, mail_pass, mail_host, content_all):
    message = MIMEMultipart()  # 實例化對象
    sender = 'husenior11123@gmail.com'
    message['From'] = Header(sender)
    message['Subject'] = Header('工資單', 'utf-8')  # 邮件主题
    receivers = 's1041026@gm.ncue.edu.tw'
    # 須將所有單獨存放的地址讀取為附件
    for receiver_name in all_data[1:]:
        path = os.path.join(os.getcwd(), "{}.docx".format(receiver_name[0]))
        # 读取附件内容
        attr = MIMEText(open(path, 'rb').read(), 'base64', 'utf-8')
        attr['Content-Type'] = 'application/octet-stream'  # 流的協議
        attr.add_header("Content-Disposition", "attachment", filename=("gbk", "", receiver_name[0] + "5月工资单.docx"))  # #解决附件未中文时，出现文件名不正确的情况
        # 添加附件到邮件
        message.attach(attr)  # 郵件附件添加完成

        msg_content = """
                   <p style="font-weight:bold;">亲爱的
                   """ + receiver_name[0] + """
                   ：</p>
                   <p style="text-indent:2em;">感谢您为公司做出的贡献！</p>
                   <p style="text-indent:2em;">本月工资已经到账，请注意查收。如未收到，请您在5个工作日内联系人力资源部。</p>
                   <p style="text-indent:2em;">您的工资单详情请查看邮件附件！</p>
                   <p style="text-indent:2em;">如有异议请联系人力资源部。</p></p>
                   <br /><br /><br />
                   <p style="text-align:right;">人力资源部<br />2021年5月31日</p>
                   """
        message.attach(MIMEText(msg_content, 'html', 'utf-8'))  # 加入郵件主內容
    try:  # 使用try-catch來捕獲發送失敗
        # 邮件协议
        smtobj = smtplib.SMTP()  # 實例化協議對象
        # 建立链接
        smtobj.connect(mail_host, 465)  # SMTP 埠號：465
        # 登陆邮箱
        smtobj.login(mail_user, mail_pass)
        smtobj.sendmail(sender, receivers, message.as_string())  # message.as_string() 加密字符串
    except smtplib.SMTPException as e:
        print('error: %s' % e)


if __name__ == '__main__':
    r_xlsx = read_excel()
    word_salary_email(r_xlsx)
    mail_host = 'ssl://smtp.gmail.com:465'  # 設置服務器
    mail_user = 'husenior11123@gmail.com'  # SMTP 帳號：[你的gmail帳號]
    mail_pass = 'wdycjpatwlswsqdo'  # SMTP 密碼：[google應用程式密碼]
    send_mail(mail_user, mail_pass, mail_host, r_xlsx)












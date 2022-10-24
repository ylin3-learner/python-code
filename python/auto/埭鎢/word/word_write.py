# coding:utf-8

# 文件生成
# Document_obj = Document()
'''
生成標題 -title_obj = DocumentObj.add_heading(標題內容, 0 <= 標題樣式等級 <= 9) -return title_obj
        內容追加: title_obj.add_run(字符串內容)
doc = Document()  # 如果是讀取文件, 則需在Document內添加文件地址; 如果為創建, 則不用

title = doc.add_heading('My title', 0)

title.add_run('\n123')
doc.save('test.docx')
========================================================
添加段落 -para_obj = document_obj.add_paragraph(段落內容)
        內容追加: para_obj.add_run(字符串內容)
        換行方式: \n
添加圖片 -image_obj = document_obj.add_picture(圖片地址, 寬, 高) -return image_obj
        寬高定義:
            from docx.shared import Inches
            add_picture(x, width=Inches(5), height=Inches(5))

            大陸的行與列與台灣相反 大陸 -橫 = 行,  縱 = 列
添加表格 -table_obj = document_obj.add_table(rows=行數, cols=列數) -return table_obj
        cell = table_obj.rows[0].cells -> cell 表格行對象
        cell[0].text = 當前行0列的內容
        cell[1].text = 當前行1列的內容
        表格追加: row_cell = table.add_row().cells() -row_cell 表格行對象
分頁 -document_obj.add_page_break()
保存生成word -doc.save(文件地址) 文件地址 -> \home\demo.docx
'''
from docx import Document
from docx.shared import Inches

doc = Document()  # 如果是讀取文件, 則需在Document內添加文件地址; 如果為創建, 則不用

title = doc.add_heading('My title', 0)

title.add_run('\n123')


p = doc.add_paragraph('歡迎來到這裡學習python')
p.add_run('\n這是關於word生成的知識')

image = doc.add_picture('logo2020.png', width=Inches(2), height=Inches(2))

titles = ['name', 'age', 'sex']
table = doc.add_table(rows=1, cols=3)

title_cells = table.rows[0].cells
title_cells[0].text = titles[0]
title_cells[1].text = titles[1]
title_cells[2].text = titles[2]

data = [
    ('xiaomu', '10', 'man'),
    ('dewei', '34', 'man'),
    ('xiaoman', '18', 'woman')
]

for d in data: # 每循環一次, 追加一行表格
    row_cells = table.add_row().cells
    print(d)
    for v in d:
        print(v)
        # row_cells[d].text = d[v]
    # print(d[0])
    # row_cells[0].text = d[0]  # name
    # row_cells[1].text = d[1]  # age
    # row_cells[2].text = d[2]  # sex

doc.add_page_break()
doc.save('test.docx')  # 保存生成word

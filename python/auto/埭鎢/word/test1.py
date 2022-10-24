# coding:utf-8

#　批量讀取文件　－Ｈｒ篩選簡歷 -利用關鍵字篩選簡歷
# 自動生成word
# 自動生成word格式
# word -> pdf
# 目標: word基本功能, 自行提升

# package - python-docx, pdfkit, pydocx

'''
python-docx -> Document

import docx
from docx import Document(文件地址) -return word文件對象(Document對象)
段落讀取 -document_obj.paragraphs -return list
使用方法: 循環獲取每個段落對象, 並調用text

表格讀取 -document_obj.tables -return list
使用方法: 通過循環, 獲取行與列的內容
獲取出有效的簡歷 -簡歷篩選
'''

from docx import Document  # Document 創建文本對象

doc = Document('文本.docx')  # doc(此包不支持格式) -> docx(支持格式)
print(doc.paragraphs)  # 返回docx.paragraphs 對象

for p in doc.paragraphs:
    print(p.text)  # text 為doc.paragraphs的屬性


for t in doc.tables:  # t 為一個個的表格對象
    for row in t.rows:  # dir(t) -讀取每一行
        _row_str = ''
        for cell in row.cells:  # 讀取行裡的小表格
            _row_str += cell.text + ','  # 獲取小表格內的字符串 -並拼接成一個字符串
        print(_row_str)  # 打印每行的信息
    print(dir(t))

for v in doc.tables:
    for column in v.columns:
        _column_str = ''
        for cell in column.cells:
            _column_str += cell.text + ','
            print(_column_str)
        print(_column_str, '***')
        print('======')

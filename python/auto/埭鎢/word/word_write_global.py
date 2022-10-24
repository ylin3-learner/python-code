# coding:utf-8

# Document_obj = Document()

# 全局樣式定義 style = Document_obj.styles['Normal'] -style == 樣式對象
'''
如何查看有哪些樣式對象: for style in Document_obj.styles:
                        print(style)
常見: -字體: style.font.name = '微軟雅黑'
     -字體顏色: style.font.color.rgb = RGBColor(255,0,0) => 調用 from docx.shared import RGBColor
     -字體大小: style.font.size = Pt(20) => 調用 from docx.shared import Pt
'''
# 文本樣式定義 -標題與段落
'''
    字體 -obj.font.name = '微軟雅黑'
    字體顏色 -obj.font.color.rgb = RGBColor(255, 0, 0)
    字體大小 -obj.font.size = Pt(20)
    標題居中 -obj.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER, LEFT, RIGHT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    字體斜體 -obj.italic = True
    字體加粗 -obj.bold = True, False -還原
'''
# 圖片的居中
'''
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER -from docx.enum.text import WD_ALIGN_PARAGRAPH
_p = p.add_run()
_p.add_picture(圖片地址)
'''
# 表格的樣式定義
'''
獲取表格樣式類型: 
獲取表格樣式: 從茫茫style海中獲取表格樣式 -from docx.enum.style import WD_STYLE_TYPE
定義表格:  document.add_table(rows=, cols=, style=)
'''
import docx
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE  # 獲取表格樣式

doc_obj = docx.Document()

for style in doc_obj.styles:
    print(style)

style = doc_obj.styles['Normal']
style.font.name = '微軟雅黑'
# style.font.color.rgb = RGBColor(255, 0, 0) -全局字體顏色
style.font.size = Pt(12)

h = doc_obj.add_heading('The first doc', 5)

h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
# h.font.size = Pt(20) -'Paragraph' object has no attribute 'font'
h.style.font.size = Pt(20)  # 標題字體大小
h.italic = True  # 但是對標題無效

_h = h.add_run('\tHahaha....')
_h.italic = True  # 對追加的內容才有斜體樣式
# 那如何讓標題本身就有斜體效果呢? 使一開始的add_heading內容為空, 標題內容由add_run()追加

_h.bold = True  # 字體加粗
print(dir(_h))
_h.underline = True  # 加底線


p = doc_obj.add_paragraph('hello world')
p.add_run('\nI use python').italic = True
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
print('========')
print(dir(WD_ALIGN_PARAGRAPH))

photo = doc_obj.add_paragraph()
photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_photo = photo.add_run()
_photo.add_picture('logo2020.png', width=Inches(5), height=Inches(2))

for i in doc_obj.styles:
    print(i.type)
    if i.type == WD_STYLE_TYPE.TABLE:  # 表格類型
        print(i.name)
t = doc_obj.add_table(rows=1, cols=2, style='Colorful Grid Accent 1')
title = ['Strength', 'Weakness']
content = ['Opportunity', 'Trick']

cell = t.rows[0].cells  # 表格列對象
cell[0].text = title[0]  # 當前列0行的內容  # 大陸的行與列與台灣相反 大陸 -橫 = 行,  縱 = 列; 台灣 -縱 = 行, 橫 = 列
cell[1].text = title[1]  # 當前列1行的內容

row_cells = t.add_row().cells
row_cells[0].text = content[0]  # Opportunity
row_cells[1].text = content[1]  # Trick

doc_obj.save('global_write.docx')
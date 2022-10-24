# coding:utf-8

# 簡歷篩選
# 已知條件: 篩選包含指定關鍵字的簡歷
# 批量讀取每一個word (glob package)
# 通過關鍵字方式篩選, 拿到目標簡歷地址

'''
from docx import Document

doc = Document('簡歷1.docx')

for p in doc.paragraphs:
    print(p.text)

# 沒有顯示, 有可能為表格類型的文檔
print('======')

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            print(cell.text)
'''

from docx import Document

import glob

class ReadDoc(object):
    def __init__(self, path):  # 初始化函數
        self.doc = Document(path)
        self.p_text = ''  # 段落文字
        self.table_text = ''  # 表格文字

        self.get_para()  # 添加進初始化函數, 使函數自動執行與獲取參數
        self.get_table()

    def get_para(self):  # 讀取段落函數
        for p in self.doc.paragraphs:
            p.text += self.p_text + '\n'

    def get_table(self):  # 讀取表格函數
        for t in self.doc.tables:
            for row in t.rows:
                _cell_str = ''
                for cell in row.cells:
                    _cell_str += cell.text + ','  # 先合成每行表格內容的文字為小字符串
                self.table_text += _cell_str + '\n'  # 合成小字符串變成表格形式

def search(path, targets):  # 關鍵字獲取函數
    res = glob.glob(path)  # 透過glob獲取當前路徑下的所有信息

    final_result = []
    is_use = True

    for i in res:
        if glob.os.path.isfile(i):  # 只判斷是否為文件
            if i.endswith('.docx'):  # 判斷是否為.docx類型文件
                # 獲取數據
                doc = ReadDoc(i)
                # 將獲取到的數據變為字符串
                p_text = doc.p_text
                t_text = doc.table_text
                # 方便進行過濾查詢
                all_text = p_text + t_text

                for target in targets:  # 所有條件都必須要滿足 -break
                    if target not in all_text:
                        is_use = False
                        break  # break：強制跳出 ❮整個❯ 迴圈

                if not is_use:  # is_use == False
                    continue  # continue：強制跳出 ❮本次❯ 迴圈，繼續進入下一圈 -因為不滿足條件
                # 滿足條件
                final_result.append(i)
    return final_result


if __name__ == '__main__':
    path = glob.os.path.join(glob.os.getcwd(), '*')  # 讀取word內所有的目錄
    res = search(path, 'Python')
    print(res)

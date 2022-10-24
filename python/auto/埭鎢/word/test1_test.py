# coding:utf-8

# 任务要求
# 1、得到文档中的总段落数
# 2、输出每个段落的内容
# 3、使用format()函数格式化输出


import docx

doc = docx.Document('文本文檔.docx')

print('段落數:', len(doc.paragraphs))
for index, _data in enumerate(doc.paragraphs):
    print('第{}段的內容是:'.format(index+1), _data.text)





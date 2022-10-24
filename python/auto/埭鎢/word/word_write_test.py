# coding:utf-8

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE

doc = Document()
style = doc.styles['Normal']
style.font.name = '微軟雅黑'
style.font.size = Pt(15)

h = doc.add_heading('慕課網是誰', 0)
p = doc.add_paragraph('我們是IT教育行業的')
# 紅色
red_p1 = p.add_run('造夢者,\t')
red_p1.font.color.rgb = RGBColor(255, 0, 0)
p.add_run('也是前研技術內容的')
# 紅色
red_p2 = p.add_run('創造者')
red_p2.font.color.rgb = RGBColor(255, 0, 0)
p.add_run('和')
# 紅色
red_p3 = p.add_run('傳播者')
red_p3.font.color.rgb = RGBColor(255, 0, 0)
p.add_run('!\n\n')

p.add_run('體系課:\n\n')
p.add_run('1.Python全棧工程師\n2.Java工程師\n3.前端工程師')

# 分頁
doc.add_page_break()
doc.add_heading('學生信息', 0)
for t in doc.styles:
    if t.type == WD_STYLE_TYPE.TABLE:
        print(t.name)
table = doc.add_table(rows=1, cols=4, style='Table Grid')  # 先寫titile
titles = ['序號', '姓名', '年齡', '身高']
data = [
    ('1', '李四','19', '167'),
    ('2','張三', '20', '174'),
]
titile_cell = table.rows[0].cells  # 表格列對象
titile_cell[0].text = titles[0]
titile_cell[1].text = titles[1]
titile_cell[2].text = titles[2]
titile_cell[3].text = titles[3]
# 在寫表格內容
for d in data:
    row_cells = table.add_row().cells
    row_cells[0].text = d[0]
    row_cells[1].text = d[1]
    row_cells[2].text = d[2]
    row_cells[3].text = d[3]

doc.save('demo.docx')
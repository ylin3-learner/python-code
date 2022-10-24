# coding:utf-8

import docx, xlrd, glob, smtplib, time
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import Header

# 打开工作薄
excel = xlrd.open_workbook("salary.xlsx")
# 读取sheet页
book = excel.sheet_by_name("研发部工资表")

"""
代码bug，1个
发送邮件bug，有多个ip时发送邮件，从第二个邮件开始会迭代第一个邮件内容和附件
"""


# 1.实现excle数据的读出
# 创建函数，读取excel表格中的内容
def read_excel():
    # 按行读取表格内容
    content_all = []
    for row_value in book.get_rows():
        content = []
        for cell_value in row_value:
            content.append(cell_value.value)
        content_all.append(content)
    return content_all


doc = docx.Document()


def word_style():
    # 设置字体和大小和格式
    style = doc.styles['Normal']
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    style.font.size = Pt(11)

    # 设置标题格式和文字大小
    title = doc.add_heading("", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _t = title.add_run("5月工资单")
    _t.font.name = "微软雅黑"
    _t.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    _t.bold = True


# 定义工资单邮件函数
def create_salarg(content_all):
    word_style()
    # 创建一个两行，X列的表格
    table = doc.add_table(rows=2, cols=book.ncols - 1, style="Light List Accent 5")

    # 设置表头
    title_cells = table.rows[0].cells

    for index, salary in enumerate(content_all[0][0:-1]):
        title_cells[index].text = salary

    # 填充单元格

    for data in content_all[1:]:
        # 设置本行单元格的数值
        # add_row:在原有表格的基础上增加一行数据
        row_cells = table.rows[1].cells
        for index, content in enumerate(data[0:-1]):
            # 数据类型转换，浮点数不能直接写入单元格，对数据进行转换
            content_to_str = str(content)
            row_cells[index].text = content_to_str
        doc.save("{}5月工资单.docx".format(data[0]))  # name
# data = [小慕	男	技术部	15000.0	1800.0	500.0	1713.6	2016.0	635.94	12934.46	xiaomu@imooc.com]

# 自定义邮件发送函数
def send_mail(mail_user, mail_pass, mail_host, content_all):
    message = MIMEMultipart()  # 定义邮件对象
    sender = '942949774@qq.com'  # 发件人邮箱
    message['From'] = Header(sender)
    message['subject'] = Header('工资单')  # 邮件主题
    for receiver_name in content_all[1:]:
        path = glob.os.path.join(glob.os.getcwd(), '{}5月工资单.docx'.format(receiver_name[0]))
        # 读取附件内容
        attr = MIMEText(open(path, 'rb').read(), 'base64', 'utf-8')

        attr['Content-Type'] = 'application/octet-stream'
        attr.add_header("Content-Disposition", "attachment", filename=("gbk", "", receiver_name[0] + "5月工资单.docx"))
        # 添加附件到邮件
        message.attach(attr)

        msg_content = """
                <p style="font-weight:bold;">亲爱的
                """ + receiver_name[0] + """
                ：</p>
                <p style="text-indent:2em;">感谢您为公司做出的贡献！</p>
                <p style="text-indent:2em;">本月工资已经到账，请注意查收。如未收到，请您在5个工作日内联系人力资源部。</p>
                <p style="text-indent:2em;">您的工资单详情请查看邮件附件！</p>
                <p style="text-indent:2em;">如有异议请联系人力资源部。</p></p>
                <br /><br /><br />
                <p style="text-align:right;">人力资源部<br />2021年5月31日</p>
                """
        message.attach(MIMEText(msg_content, 'html', 'utf-8'))

        try:
            # 邮件协议
            smtobj = smtplib.SMTP()
            # 建立链接
            smtobj.connect(mail_host, 587)
            # 登陆邮箱
            smtobj.login(mail_user, mail_pass)
            smtobj.sendmail(sender, ['qb_wangyi@163.com'], message.as_string())
        except smtplib.SMTPException as e:
            print("error: %s" % e)


if __name__ == '__main__':
    # path = glob.os.path.join(glob.os.getcwd(), 'salary.xlsx')
    resp = read_excel()
    create_salarg(resp)
    # 设置服务器
    mail_host = 'smtp.qq.com'
    # 用户登录名
    # mail_user = 'qb_wangyi'
    mail_user = '942949774'
    # 设置个人口令
    mail_pass = 'turyximpqzdwbbhj'
    # mail_pass = 'EUEVGMXJGGFCODPE'
    send_mail(mail_user, mail_pass, mail_host, resp)

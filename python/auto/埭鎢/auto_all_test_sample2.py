# coding: utf-8

import xlrd
import glob

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

import smtplib
from email.mime.text import MIMEText
from email.header import Header
from email.mime.multipart import MIMEMultipart

import multiprocessing


def read_data(path):
    excel_obj = xlrd.open_workbook(path)
    book_obj = excel_obj.sheet_by_name('研发部工资表')
    datas = []
    for row in book_obj.get_rows():
        content = [j.value for j in row]
        datas.append(content)
    return datas


def write_word(datas, q):
    doc = Document()
    # 设置字体和大小
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 中文字体
    style.font.size = Pt(6)

    # 设置word的标题样式
    title = doc.add_heading('', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    t = title.add_run('5月工资单')
    t.font.name = '微软雅黑'
    t._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 中文字体
    t.bold = True

    # 插入表格
    rows = len(datas)
    cols = len(datas[0])
    table = doc.add_table(rows, cols, style='Medium List 1 Accent 1')

    for i in range(rows):
        cells = table.rows[i].cells
        for j in range(cols):
            cells[j].text = str(datas[i][j])
    filename = f'{datas[1][0]}_5月工资单.docx'
    doc.save(filename)
    q.put(filename)
    # return filename


def send_mail(name, email, q):
    mail_host = 'smtp.qq.com'
    mail_user = '45xxxx65@qq.com'
    mail_pass = 'lqdkcccmwmnjcaad'

    sender = '45xxxx65@qq.com'
    receivers = email

    message = MIMEMultipart()

    message['From'] = Header(sender)

    filename = q.get()
    subject = filename.split('.')[0].split('_')[-1]
    message['Subject'] = Header(subject, 'utf-8')

    attr = MIMEText(open(filename, 'rb').read(), 'base64', 'utf-8')
    attr['Content-Type'] = 'application/octet-stream'
    attr.add_header("Content-Disposition", "attachment", filename=("gbk", "", filename))
    message.attach(attr)
    msg_content = """
    <p style="font-weight:bold;">亲爱的
    """ + name + """
    ：</p>
    <p style="text-indent:2em;">感谢您为公司做出的贡献！</p>
    <p style="text-indent:2em;">本月工资已经到账，请注意查收。如未收到，请您在5个工作日内联系人力资源部。</p>
    <p style="text-indent:2em;">您的工资单详情请查看邮件附件！</p>
    <p style="text-indent:2em;">如有异议请联系人力资源部。</p></p>
    <br /><br /><br />
    <p style="text-align:right;">人力资源部<br />2021年5月31日</p>
    """
    message.attach(MIMEText(msg_content, 'html', 'utf-8'))
    try:
        smtpobj = smtplib.SMTP()
        smtpobj.connect(mail_host, 25)
        smtpobj.login(mail_user, mail_pass)
        smtpobj.sendmail(sender, receivers, message.as_string())
    except smtpobj.SMTPException as e:
        print('error: %s' % e)


if __name__ == '__main__':

    path = glob.os.path.join(glob.os.getcwd(), 'salary.xlsx')
    data = read_data(path)

    recs = [(data[0], data[i]) for i in range(1, len(data))]

    q = multiprocessing.Queue()
    # 为了写word，每个元素里面第一个元素是表头
    for datas in recs:
        p_w = multiprocessing.Process(target=write_word, args=(datas, q,))
        p_mail = multiprocessing.Process(target=send_mail, args=(datas[1][0], datas[1][-1], q,))
        for p in (p_w, p_mail):
            p.start()
        for p in (p_w, p_mail):
            p.join()